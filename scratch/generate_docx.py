"""
generate_docx.py  – Membuat file .docx native menggunakan python-docx
Struktur per halaman:
  Page 1  : Cover (Halaman Judul)
  Page 2  : Lembar Pengesahan
  Page 3  : ABSTRACT (English)
  Page 4  : ABSTRAK (Indonesian)
  Page 5  : Kata Pengantar
  Page 6+ : BAB I, BAB II, BAB III, BAB IV, BAB V
"""

import re, base64, io, ssl, urllib.request
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

SRC_FILE      = r'C:\Users\Haikal Al-Ghifari\.gemini\antigravity-ide\brain\0e56245d-303e-4e0e-a812-ea8f17fa3d46\Laporan_Lengkap.md'
OUT_FILE      = r'd:\Data Joki\ComproRRK\public\Laporan_Kerja_Praktek.docx'
LOGO_RRK_FILE = r'd:\Data Joki\ComproRRK\public\logo.png'
LOGO_UNIKOM   = r'd:\Data Joki\ComproRRK\public\logo_unikom.png'
SCREENSHOT_LOGIN_FILE     = r'C:\Users\Haikal Al-Ghifari\.gemini\antigravity-ide\brain\4987b58e-f65c-4ba7-9cfe-908757d8bafe\screenshot_login.png'
SCREENSHOT_DASHBOARD_FILE = r'C:\Users\Haikal Al-Ghifari\.gemini\antigravity-ide\brain\4987b58e-f65c-4ba7-9cfe-908757d8bafe\screenshot_dashboard.png'
SCREENSHOT_LANDING_ID_FILE = r'C:\Users\Haikal Al-Ghifari\.gemini\antigravity-ide\brain\4987b58e-f65c-4ba7-9cfe-908757d8bafe\screenshot_landing_id.png'
SCREENSHOT_LANDING_EN_FILE = r'C:\Users\Haikal Al-Ghifari\.gemini\antigravity-ide\brain\4987b58e-f65c-4ba7-9cfe-908757d8bafe\screenshot_landing_en.png'

MAX_IMG_W_CM = 14.0   # usable page width = 21cm - 4cm left - 3cm right

ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode    = ssl.CERT_NONE
HEADERS = {'User-Agent': 'Mozilla/5.0'}


# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────
import time

def fetch_image(url):
    url = url.replace('&amp;','&')
    max_retries = 4
    for attempt in range(1, max_retries + 1):
        try:
            req = urllib.request.Request(url, headers=HEADERS)
            with urllib.request.urlopen(req, timeout=20, context=ctx) as r:
                data = r.read()
            img = Image.open(io.BytesIO(data)).convert('RGB')
            print(f"  OK {url[:55]}")
            return img
        except Exception as e:
            if attempt == max_retries:
                print(f"  FAIL {url[:55]}: {e}")
                # Fallback: try without scale=3 if it failed on scale=3
                if 'scale=' in url:
                    fallback_url = re.sub(r'[?&]scale=\d+', '', url)
                    print(f"  RETRACTING to scale=1: {fallback_url[:55]}")
                    try:
                        req_f = urllib.request.Request(fallback_url, headers=HEADERS)
                        with urllib.request.urlopen(req_f, timeout=20, context=ctx) as r_f:
                            data_f = r_f.read()
                        img_f = Image.open(io.BytesIO(data_f)).convert('RGB')
                        # Override DPI flag to 96 since scale is 1
                        img_f.info['DPI'] = 96
                        return img_f
                    except Exception as ef:
                        print(f"  FALLBACK FAIL: {ef}")
                return None
            else:
                # Wait before retrying
                time.sleep(2)


def render_mermaid(code):
    if '%%{init' not in code:
        init = ("%%{init: {'theme':'neutral','themeVariables':{"
                "'background':'#ffffff','primaryColor':'#ffffff',"
                "'primaryTextColor':'#000000','primaryBorderColor':'#000000',"
                "'lineColor':'#000000','secondaryColor':'#ffffff','tertiaryColor':'#ffffff',"
                "'actorBkg':'#ffffff','actorBorder':'#000000','actorTextColor':'#000000',"
                "'noteBkgColor':'#ffffff','noteBorderColor':'#000000'}}}%%\n")
        code = init + code
    code = re.sub(r'fill:#(?!ffffff)[0-9a-fA-F]{6}', 'fill:#ffffff', code)
    encoded = base64.urlsafe_b64encode(code.encode()).decode().rstrip('=')
    url = f'https://mermaid.ink/img/{encoded}'
    return fetch_image(url)


def set_margins(doc, top=4, bottom=3, left=4, right=3):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def set_font(run, size=12, bold=False, italic=False, underline=False):
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.color.rgb = RGBColor(0,0,0)


def add_para(doc, text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
             italic=False, size=12, space_after=6, space_before=0, first_indent=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    if text:
        # Handle inline **bold** and *italic* by splitting
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                set_font(r, size=size, bold=True)
            elif part.startswith('*') and part.endswith('*'):
                r = p.add_run(part[1:-1])
                set_font(r, size=size, italic=True)
            else:
                # Strip remaining markdown markers
                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', part)
                clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
                clean = re.sub(r'`([^`]+)`', r'\1', clean)
                clean = clean.replace('&nbsp;', ' ').replace('&amp;', '&')
                if clean:
                    r = p.add_run(clean)
                    set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1, center=True, size=None, uppercase=False):
    if size is None:
        size = 14 if level == 1 else 12
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    text_out = text.upper() if uppercase else text
    
    # Use native Word heading style so the Table of Contents (TOC) builder detects it
    p = doc.add_paragraph(style=f'Heading {level}')
    p.alignment = align
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.keep_with_next = (level > 1)
    if level == 1:
        no_break_headings = {
            'PENDAHULUAN',
            'TINJAUAN PUSTAKA',
            'PEMBAHASAN',
            'IMPLEMENTASI DAN PENGUJIAN',
            'PENUTUP'
        }
        if text.strip().upper() not in no_break_headings:
            p.paragraph_format.page_break_before = True
    
    # Override Heading style appearance to match layout requirements
    run = p.add_run(text_out)
    set_font(run, size=size, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    return p


def add_image_from_pil(doc, img, max_width_cm=MAX_IMG_W_CM, max_height_cm=7.5, caption_text=None):
    buf = io.BytesIO()
    
    # Detect if scale=1 fallback was used (DPI=96), otherwise use default 288
    DPI = img.info.get('DPI', 288) if (hasattr(img, 'info') and img.info) else 288
    
    # Calculate target dimensions in pixels
    target_w_px = max_width_cm * DPI / 2.54
    target_h_px = max_height_cm * DPI / 2.54
    
    # Calculate scale factor to fit within both dimensions
    scale_w = target_w_px / img.width
    scale_h = target_h_px / img.height
    scale = min(scale_w, scale_h)
    
    # Scale to fit targets (either down or up) for maximum readability
    new_w = max(1, int(img.width * scale))
    new_h = max(1, int(img.height * scale))
    img = img.resize((new_w, new_h), Image.LANCZOS)
    print(f"  IMAGE RESIZE: orig={img.width/scale:.0f}x{img.height/scale:.0f} | scale={scale:.2f} | target={max_width_cm}x{max_height_cm}cm | final={new_w}x{new_h}px ({new_w*2.54/DPI:.2f}x{new_h*2.54/DPI:.2f}cm)")
        
    img.save(buf, format='PNG')
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(buf, width=Cm(img.width * 2.54 / DPI), height=Cm(img.height * 2.54 / DPI))
    
    if caption_text:
        add_caption(doc, caption_text, is_table=False)
    return p


def add_image_from_file(doc, path, width_cm=5, caption_text=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    try:
        with Image.open(path) as img:
            w_px, h_px = img.size
        ratio = h_px / w_px
        run.add_picture(path, width=Cm(width_cm), height=Cm(width_cm * ratio))
    except Exception as e:
        print(f"  add_image_from_file err: {e}")
        run.add_picture(path, width=Cm(width_cm))
        
    if caption_text:
        add_caption(doc, caption_text, is_table=False)
    return p


def page_break(doc):
    """Add a page break. Reuses the last paragraph if it is empty,
    avoiding the blank-page artefact caused by a spacer paragraph
    landing at the top of the next page."""
    from docx.enum.text import WD_BREAK
    paragraphs = doc.paragraphs
    # If the most-recent paragraph is blank, put the break there
    if paragraphs and not paragraphs[-1].text.strip():
        last_p = paragraphs[-1]
        last_p.paragraph_format.space_after  = Pt(0)
        last_p.paragraph_format.space_before = Pt(0)
        last_p.add_run().add_break(WD_BREAK.PAGE)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc, field_code="TOC \\o \"1-3\" \\h \\z \\u"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), field_code)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    r_font = OxmlElement('w:rFonts')
    r_font.set(qn('w:ascii'), 'Times New Roman')
    r_font.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(r_font)
    
    r_size = OxmlElement('w:sz')
    r_size.set(qn('w:val'), '24')  # 12pt
    rPr.append(r_size)
    
    r_italic = OxmlElement('w:i')
    rPr.append(r_italic)
    
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = "(Klik kanan di sini dan pilih 'Update Field' untuk memperbarui daftar secara otomatis)"
    r.append(t)
    fldSimple.append(r)
    
    p._p.append(fldSimple)
    return p



def configure_section_page_numbering(section, fmt='decimal', start_val=None):
    """Configure page numbering format and starting value on a section properties XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        # Strict schema order insertion (must be before w:cols or w:docGrid)
        inserted = False
        for child in sectPr:
            if child.tag in (qn('w:cols'), qn('w:docGrid'), qn('w:vAlign'), qn('w:formProt')):
                child.addprevious(pgNumType)
                inserted = True
                break
        if not inserted:
            sectPr.append(pgNumType)
    
    pgNumType.set(qn('w:fmt'), fmt)
    if start_val is not None:
        pgNumType.set(qn('w:start'), str(start_val))
    else:
        # Remove starting value attribute to continue numbering from previous section
        pgNumType.attrib.pop(qn('w:start'), None)


def clean_text(text):
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()


def add_list_item(doc, text, ordered=False, num=None, letter=None, size=12):
    """Rule 4: lists use numbers (ordered) or letters a./b./c. (unordered), NOT bullets."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent   = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.64)
    if ordered and num:
        prefix = f'{num}. '
    elif letter:
        prefix = f'{letter}. '
    else:
        prefix = 'a. '
    r = p.add_run(prefix)
    set_font(r, size=size)
    # Handle inline bold/italic in list items
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_font(r, size=size, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            set_font(r, size=size, italic=True)
        else:
            clean = clean_text(part)
            if clean:
                r = p.add_run(clean)
                set_font(r, size=size)
    return p


def add_table_from_md(doc, header_cells, rows, borderless=False):
    """
    Rule 5: single-spacing cells.
    Plus: full-page-width table, narrow 'No' column, proportional others.
    """
    cols = len(header_cells)
    if cols == 0:
        return None

    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    tbl.style     = 'Normal Table' if borderless else 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit   = False

    PAGE_W_CM = 14.0
    NO_COLS   = {'no', 'no.', '#', 'nomor'}
    first_is_no = cols > 1 and clean_text(header_cells[0]).strip().lower() in NO_COLS

    if first_is_no:
        no_w        = 1.2
        other_w     = (PAGE_W_CM - no_w) / (cols - 1)
        col_widths  = [no_w] + [other_w] * (cols - 1)
    else:
        col_widths  = [PAGE_W_CM / cols] * cols

    # Apply widths using built-in cell.width (which handles XML schema order safely)
    for ci, w_cm in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[ci].width = Cm(w_cm)

    # ── Cell paragraph formatter ──────────────────────────────────────────────
    def set_cell_para(para, bold=False, center=False):
        para.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                          if center else WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_after        = Pt(0)
        para.paragraph_format.space_before       = Pt(0)
        para.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.SINGLE
        for run in para.runs:
            set_font(run, size=11, bold=bold)

    # ── Header row ────────────────────────────────────────────────────────────
    hdr_row = tbl.rows[0]
    for i, cell_text in enumerate(header_cells):
        cell = hdr_row.cells[i]
        cell.text = clean_text(cell_text)
        set_cell_para(cell.paragraphs[0], bold=True, center=True)

    # ── Data rows ─────────────────────────────────────────────────────────────
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, cell_text in enumerate(row_data[:cols]):
            cell = row.cells[ci]
            cell.text = clean_text(cell_text)
            is_no_col = first_is_no and ci == 0
            set_cell_para(cell.paragraphs[0], bold=False, center=is_no_col)

    return tbl



def add_caption(doc, text, is_table=True):
    """Rule 6: table title above table / figure caption below image.
       12pt, Times New Roman, single spacing, centered, style='Caption'."""
    p = doc.add_paragraph(style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_table:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
    else:
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.keep_with_next = True
    
    # Split caption: bold label (e.g. "Tabel 3.1") + rest of text
    m = re.match(r'^(\*{0,2}(Tabel|Gambar|Diagram|Grafik|Bagan)\s+(\d+)\.(\d+)\*{0,2})(.*)', text.strip(), flags=re.IGNORECASE)
    if m:
        type_word = m.group(2).strip().capitalize()
        if type_word in ['Diagram', 'Grafik', 'Bagan']:
            type_word = 'Gambar'
        
        ch_num = m.group(3)
        idx = m.group(4)
        rest  = m.group(5).strip()
        
        r1 = p.add_run(f'{type_word} {ch_num}.')
        set_font(r1, size=12, bold=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        r_field = p.add_run()
        set_font(r_field, size=12, bold=True)
        r_field.font.color.rgb = RGBColor(0, 0, 0)
        
        instruction = f' SEQ {type_word} \\* ARABIC '
        if str(idx) == '1':
            instruction += '\\r 1 '
            
        def add_seq_field(run, instr, show_txt):
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = instr
            run._r.append(instrText)

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            run._r.append(fldChar2)

            t = OxmlElement('w:t')
            t.text = show_txt
            run._r.append(t)

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar3)
            
        add_seq_field(r_field, instruction, str(idx))
        
        if rest:
            r2 = p.add_run(f' {rest}')
            set_font(r2, size=12)
            r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        # Fallback for old format
        m_old = re.match(r'^(\*{0,2}(?:Tabel|Gambar|Diagram|Grafik|Bagan)\s+[\d.]+\*{0,2})(.*)', text.strip(), flags=re.IGNORECASE)
        if m_old:
            label = m_old.group(1).strip('*').strip()
            rest  = m_old.group(2).strip()
            r = p.add_run(label)
            set_font(r, size=12, bold=True)
            r.font.color.rgb = RGBColor(0, 0, 0)
            if rest:
                r2 = p.add_run(f' {rest}')
                set_font(r2, size=12)
                r2.font.color.rgb = RGBColor(0, 0, 0)
        else:
            r = p.add_run(clean_text(text))
            set_font(r, size=12, bold=True)
            r.font.color.rgb = RGBColor(0, 0, 0)
    return p


# ─────────────────────────────────────────────────────────────
# Load + pre-render mermaid diagrams
# ─────────────────────────────────────────────────────────────
print("Loading source...")
with open(SRC_FILE, encoding='utf-8') as f:
    src = f.read()

print("Rendering mermaid diagrams...")
mermaid_images = []
def capture_mermaid(m):
    code = m.group(1).strip()
    img = render_mermaid(code)
    idx = len(mermaid_images)
    mermaid_images.append(img)
    return f'[[MERMAID:{idx}]]'

src = re.sub(r'```mermaid(.*?)```', capture_mermaid, src, flags=re.DOTALL)

# ─────────────────────────────────────────────────────────────
# Split source into sections by page-break divs
# ─────────────────────────────────────────────────────────────
# Remove <style> block
src = re.sub(r'<style[^>]*>.*?</style>', '', src, flags=re.DOTALL | re.IGNORECASE)
# Remove div tags (but keep content)
src = re.sub(r'<div[^>]*page-break-after:[^>]*>\s*</div>', '[[PAGEBREAK]]', src, flags=re.IGNORECASE)
src = re.sub(r'<div[^>]*page-break-before:[^>]*>', '[[PAGEBREAK]]', src, flags=re.IGNORECASE)
src = re.sub(r'</?div[^>]*>', '', src, flags=re.IGNORECASE)

sections = src.split('[[PAGEBREAK]]')
print(f"Found {len(sections)} sections.")


# ─────────────────────────────────────────────────────────────
# Process a single section's text content → add to doc
# ─────────────────────────────────────────────────────────────
def process_section(doc, section_text, center_default=False):
    lines = section_text.strip().splitlines()
    i = 0
    # Accumulate table lines
    table_buffer = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty
        if not stripped:
            i += 1
            continue
        
        # Mermaid image placeholder
        mm = re.match(r'\[\[MERMAID:(\d+)\]\]', stripped)
        if mm:
            idx = int(mm.group(1))
            img = mermaid_images[idx] if idx < len(mermaid_images) else None
            if img:
                add_image_from_pil(doc, img)
            i += 1; continue
        
        # Markdown heading
        hm = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if hm:
            level = len(hm.group(1))
            text  = hm.group(2)
            center = (level == 1 or center_default)
            size = 14 if level == 1 else 12
            add_heading(doc, text, level=level, center=center, size=size)
            i += 1; continue
        
        # HTML table (passthrough)
        if stripped.startswith('<table') or stripped.startswith('<TABLE'):
            # Collect until </table>
            tbl_lines = []
            while i < len(lines) and not re.search(r'</table>', lines[i], re.IGNORECASE):
                tbl_lines.append(lines[i])
                i += 1
            if i < len(lines):
                tbl_lines.append(lines[i])
                i += 1
            # Parse simple HTML table
            raw = '\n'.join(tbl_lines)
            # Extract cells (basic)
            rows_html = re.findall(r'<tr[^>]*>(.*?)</tr>', raw, re.DOTALL | re.IGNORECASE)
            if rows_html:
                doc_rows = []
                header = []
                for ri, row_html in enumerate(rows_html):
                    cells_td = re.findall(r'<td[^>]*>(.*?)</td>', row_html, re.DOTALL | re.IGNORECASE)
                    cells_th = re.findall(r'<th[^>]*>(.*?)</th>', row_html, re.DOTALL | re.IGNORECASE)
                    if cells_th and ri == 0:
                        header = [clean_text(c) for c in cells_th]
                    elif cells_td:
                        doc_rows.append([clean_text(c) for c in cells_td])
                    elif cells_th:
                        doc_rows.append([clean_text(c) for c in cells_th])
                if doc_rows and not header:
                    header = doc_rows.pop(0)
                if header:
                    add_table_from_md(doc, header, doc_rows)
                    doc.add_paragraph()  # spacer after table
            continue
        
        # Markdown table (pipe-delimited)
        if stripped.startswith('|'):
            table_buffer.append(stripped)
            i += 1
            # Collect all table lines
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_buffer.append(lines[i].strip())
                i += 1
            # Parse
            header_row = None
            data_rows  = []
            for tl in table_buffer:
                if re.match(r'^\|?[\-: |]+\|?$', tl):
                    continue
                cells = [c.strip() for c in re.split(r'(?<!\\)\|', tl.strip('|'))]
                if header_row is None:
                    header_row = cells
                else:
                    data_rows.append(cells)
            if header_row:
                add_table_from_md(doc, header_row, data_rows)
                doc.add_paragraph()
            table_buffer = []
            continue
        
        # Ordered list item
        ol_m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if ol_m:
            num  = int(ol_m.group(1))
            text = ol_m.group(2)
            add_list_item(doc, text, ordered=True, num=num)
            i += 1; continue
        
        # Unordered list item
        ul_m = re.match(r'^[-*+]\s+(.+)$', stripped)
        if ul_m:
            text = ul_m.group(1)
            add_list_item(doc, text, ordered=False)
            i += 1; continue
        
        # Skip raw HTML tags (br, img already embedded, etc.)
        if re.match(r'^<[a-zA-Z/!]', stripped) and '[[MERMAID' not in stripped:
            i += 1; continue
        
        # Regular paragraph
        plain = clean_text(stripped)
        if plain:
            align = WD_ALIGN_PARAGRAPH.CENTER if center_default else WD_ALIGN_PARAGRAPH.JUSTIFY
            first_indent = not center_default
            add_para(doc, plain, align=align, first_indent=first_indent)
        i += 1


# ─────────────────────────────────────────────────────────────
# Build the document
# ─────────────────────────────────────────────────────────────
doc = Document()
set_margins(doc, top=4, bottom=3, left=4, right=3)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

section_map = [
    'cover',
    'lembar_pengesahan',
    'abstract',
    'abstrak',
    'kata_pengantar',
]
# Remaining sections are BAB I, BAB II, etc.

# Initialize figure and table counters per chapter
img_counters = {0: 0, 1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
tbl_counters = {0: 0, 1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
current_chapter = 0

for si, section_text in enumerate(sections):
    sec_name = section_map[si] if si < len(section_map) else f'bab_{si - len(section_map) + 1}'
    
    if si > 0:
        from docx.enum.section import WD_ORIENT, WD_SECTION
        last_sec = doc.sections[-1]
        
        # Create a new section for the prefaces or if reverting from landscape/narrow
        is_preface_start = (sec_name == 'lembar_pengesahan')
        is_revert_section = (last_sec.orientation == WD_ORIENT.LANDSCAPE or last_sec.left_margin < Cm(3.0))
        
        if is_preface_start or is_revert_section:
            new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
            new_sec.orientation = WD_ORIENT.PORTRAIT
            w, h = new_sec.page_width, new_sec.page_height
            new_sec.page_width = min(w, h)
            new_sec.page_height = max(w, h)
            new_sec.top_margin = Cm(4)
            new_sec.bottom_margin = Cm(3)
            new_sec.left_margin = Cm(4)
            new_sec.right_margin = Cm(3)
            
            if is_preface_start:
                # Format Section 1 (Prefaces) as lowercase Roman starting at ii (cover is i)
                configure_section_page_numbering(new_sec, fmt='romanLower', start_val=2)
        else:
            pass
    
    # Update current chapter based on section name
    if sec_name.startswith('bab_'):
        current_chapter = int(sec_name.split('_')[1])
        
    # Programmatically insert DAFTAR ISI, DAFTAR GAMBAR, DAFTAR TABEL right before BAB I
    if sec_name == 'bab_1':
        # 1. DAFTAR ISI
        add_heading(doc, 'DAFTAR ISI', level=1, center=True, size=14)
        doc.add_paragraph()
        add_toc(doc, field_code='TOC \\o "1-3" \\h \\z \\u')
        
        # 2. DAFTAR GAMBAR
        add_heading(doc, 'DAFTAR GAMBAR', level=1, center=True, size=14)
        doc.add_paragraph()
        add_toc(doc, field_code='TOC \\h \\z \\c "Gambar"')
        
        # 3. DAFTAR TABEL
        add_heading(doc, 'DAFTAR TABEL', level=1, center=True, size=14)
        doc.add_paragraph()
        add_toc(doc, field_code='TOC \\h \\z \\c "Tabel"')
        
        # Now transition to standard Arabic page numbering starting at 1 for BAB I
        from docx.enum.section import WD_SECTION, WD_ORIENT
        bab1_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        bab1_sec.orientation = WD_ORIENT.PORTRAIT
        w, h = bab1_sec.page_width, bab1_sec.page_height
        bab1_sec.page_width = min(w, h)
        bab1_sec.page_height = max(w, h)
        bab1_sec.top_margin = Cm(4)
        bab1_sec.bottom_margin = Cm(3)
        bab1_sec.left_margin = Cm(4)
        bab1_sec.right_margin = Cm(3)
        
        configure_section_page_numbering(bab1_sec, fmt='decimal', start_val=None)
    # center_heading: TRUE means h1 headings in this section are centered
    # center_body:    TRUE means regular paragraphs are also centered (only for COVER)
    center_heading  = sec_name in ('cover', 'lembar_pengesahan', 'abstract', 'abstrak')
    center_body     = sec_name == 'cover'   # only cover has fully-centered body text
    
    print(f"Processing section {si+1}: {sec_name}")

    # Reset unordered-list letter counter for each new section
    ul_letter_idx = 0
    next_para_no_indent = False
    last_level = 0
    last_text = ""

    # ── SPECIAL: build cover page manually (exact reference format) ───────────
    if sec_name == 'cover':
        C = WD_ALIGN_PARAGRAPH.CENTER

        def cp(txt='', bold=False, size=12, space_after=4, italic=False):
            p = doc.add_paragraph()
            p.alignment = C
            p.paragraph_format.space_after  = Pt(space_after)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if txt:
                r = p.add_run(txt)
                set_font(r, size=size, bold=bold, italic=italic)
            return p

        # Title (wrapped exactly like in the screenshot)
        cp('PERANCANGAN DAN IMPLEMENTASI APLIKASI COMPANY\nPROFILE DAN SISTEM MANAJEMEN KONTEN (CMS)\nBERBASIS WEB PADA PT RIZKY RIJAYA KARYA',
           bold=True, size=14, space_after=12)

        # KERJA PRAKTEK
        cp('KERJA PRAKTEK', bold=True, size=12, space_after=6)

        # Submission text
        cp('Diajukan Untuk Memenuhi Tugas Mata Kuliah Kerja Praktek', size=12, space_after=18)

        # Oleh
        cp('Oleh:', size=12, space_after=4)
        cp('1012079  Mochammad Alif Firmansyah', bold=True, size=12, space_after=18)

        # UNIKOM Logo (transparent PNG, centered with proper spacing)
        def cover_img(path, width_cm=5):
            p = doc.add_paragraph()
            p.alignment = C
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(18)
            try:
                with Image.open(path) as img_logo:
                    w_px, h_px = img_logo.size
                ratio = h_px / w_px
                p.add_run().add_picture(path, width=Cm(width_cm), height=Cm(width_cm * ratio))
            except Exception as e:
                print(f"  img err: {e}")
                try:
                    p.add_run().add_picture(path, width=Cm(width_cm))
                except:
                    pass
            return p

        cover_img(LOGO_UNIKOM, width_cm=4.2)

        # Pembimbing
        cp('Pembimbing:', size=12, space_after=4)
        cp('Iskandar Ikbal, S.T., M.Kom.', bold=True, size=12, space_after=32)

        # Department info
        cp('JURUSAN TEKNIK INFORMATIKA', bold=True, size=12, space_after=0)
        cp('FAKULTAS TEKNIK DAN ILMU KOMPUTER', bold=True, size=12, space_after=0)
        cp('UNIVERSITAS KOMPUTER INDONESIA', bold=True, size=12, space_after=0)
        cp('BANDUNG 2026', bold=True, size=12, space_after=0)

        continue   # skip normal section processing for cover

    # ── SPECIAL: build Lembar Pengesahan manually (exact reference format) ────
    if sec_name == 'lembar_pengesahan':
        C = WD_ALIGN_PARAGRAPH.CENTER

        def lp(txt='', bold=False, size=12, space_after=4, underline=False, line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE):
            p = doc.add_paragraph()
            p.alignment = C
            p.paragraph_format.space_after  = Pt(space_after)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing_rule = line_spacing
            if txt:
                r = p.add_run(txt)
                set_font(r, size=size, bold=bold, underline=underline)
            return p

        # Title
        lp('LEMBAR PENGESAHAN', bold=True, size=14, space_after=12)
        lp('LAPORAN HASIL KERJA PRAKTEK', bold=True, size=12, space_after=18)

        # Report Title
        lp('“PERANCANGAN DAN IMPLEMENTASI APLIKASI COMPANY PROFILE\nDAN SISTEM MANAJEMEN KONTEN (CMS) BERBASIS WEB\nPADA PT RIZKY RIJAYA KARYA”',
           bold=True, size=12, space_after=18)

        # Student Info (Name and NIM side-by-side spaced)
        lp('MOCHAMMAD ALIF FIRMANSYAH      1012079', bold=True, size=12, space_after=18)

        # Date
        lp('Bandung, 18 Juli 2026', size=12, space_after=18)

        # Signatures table (borderless)
        tbl = doc.add_table(rows=3, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        # Set column widths (usable width = 14 cm, so 7 cm each)
        for row in tbl.rows:
            row.cells[0].width = Cm(7)
            row.cells[1].width = Cm(7)

        # Add signature labels (Row 0)
        p1 = tbl.rows[0].cells[0].paragraphs[0]
        p1.alignment = C
        p1.paragraph_format.space_after = Pt(0)
        set_font(p1.add_run('Pembimbing Lapangan'), size=12)

        p2 = tbl.rows[0].cells[1].paragraphs[0]
        p2.alignment = C
        p2.paragraph_format.space_after = Pt(0)
        set_font(p2.add_run('Pembimbing Akademik'), size=12)

        # Spacer row (Row 1) - signature height
        tbl.rows[1].height = Cm(2.2)
        tbl.rows[1].cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
        tbl.rows[1].cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)

        # Names row (Row 2) - Underlined names
        p3 = tbl.rows[2].cells[0].paragraphs[0]
        p3.alignment = C
        p3.paragraph_format.space_after = Pt(0)
        p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        set_font(p3.add_run('Muhammad Ilham Ziarman'), size=12, bold=True, underline=True)
        # NIP underneath Muhammad Ilham Ziarman
        p3_nip = tbl.rows[2].cells[0].add_paragraph()
        p3_nip.alignment = C
        p3_nip.paragraph_format.space_before = Pt(0)
        p3_nip.paragraph_format.space_after = Pt(0)
        set_font(p3_nip.add_run('NIP: 1909006'), size=12)

        p4 = tbl.rows[2].cells[1].paragraphs[0]
        p4.alignment = C
        p4.paragraph_format.space_after = Pt(0)
        p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        set_font(p4.add_run('Iskandar Ikbal, S.T., M.Kom.'), size=12, bold=True, underline=True)
        # NIP & NIDN underneath Iskandar Ikbal (on separate lines)
        p4_nip = tbl.rows[2].cells[1].add_paragraph()
        p4_nip.alignment = C
        p4_nip.paragraph_format.space_before = Pt(0)
        p4_nip.paragraph_format.space_after = Pt(0)
        set_font(p4_nip.add_run('NIP: 41277006020'), size=12)
        
        p4_nidn = tbl.rows[2].cells[1].add_paragraph()
        p4_nidn.alignment = C
        p4_nidn.paragraph_format.space_before = Pt(0)
        p4_nidn.paragraph_format.space_after = Pt(0)
        set_font(p4_nidn.add_run('NIDN: 0408078002'), size=12)

        # Spacer before Mengetahui
        lp(space_after=18)

        # Mengetahui section
        lp('Mengetahui,', size=12, space_after=0)
        lp('Ketua Program Studi Teknik Informatika', size=12, space_after=0)

        # Spacer for Kaprodi signature
        p_kaprodi_sig = doc.add_paragraph()
        p_kaprodi_sig.paragraph_format.space_before = Pt(0)
        p_kaprodi_sig.paragraph_format.space_after = Pt(0)
        r_sig = p_kaprodi_sig.add_run()
        set_font(r_sig, size=12)
        p_kaprodi_sig.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p_kaprodi_sig.paragraph_format.line_spacing = 3.5

        # Kaprodi Name (Underlined)
        lp('Dedeng Hirawan S.Kom., M.Kom.', bold=True, size=12, space_after=0, underline=True)
        lp('NIP: 4127 70 06 135', size=12, space_after=0)
        lp('NIDN: 0411068810', size=12, space_after=0)

        continue   # skip normal section processing for lembar_pengesahan

    # ── Fix local logo reference ──────────────────────────────────────────────
    section_text = re.sub(
        r'(?:src=["\'])?(?:/public/logo[^"\'<\s]*|/logo\.png)["\']?',
        '[[LOCAL_LOGO]]', section_text
    )
    # Remove broken UNIKOM img from any other section (it's manually added on cover)
    section_text = re.sub(r'<img[^>]+Logo_Unikom[^>]*/>', '', section_text, flags=re.IGNORECASE)

    lines = section_text.strip().splitlines()
    i = 0
    table_buffer = []

    
    while i < len(lines):
        line  = lines[i]
        stripped = line.strip()
        
        if not stripped:
            i += 1; continue
        
        # Local RRK logo (embedded in body sections like 2.1.2)
        if '[[LOCAL_LOGO]]' in stripped:
            try:
                add_image_from_file(doc, LOGO_RRK_FILE, width_cm=5)
            except Exception as e:
                print(f"  Logo err: {e}")
            i += 1; continue
            
        # Screenshots in BAB IV
        if '*(Masukkan Screenshot Halaman Login)*' in stripped:
            try:
                ch_num = 4
                img_counters[ch_num] += 1
                caption = f"Gambar {ch_num}.{img_counters[ch_num]} Antarmuka Halaman Login"
                add_image_from_file(doc, SCREENSHOT_LOGIN_FILE, width_cm=14.0, caption_text=caption)
            except Exception as e:
                print(f"  Login screenshot err: {e}")
            i += 1; continue

        if '*(Masukkan Screenshot Dashboard Admin)*' in stripped:
            try:
                ch_num = 4
                base_dir = r"C:\Users\Haikal Al-Ghifari\.gemini\antigravity-ide\brain\4987b58e-f65c-4ba7-9cfe-908757d8bafe"
                
                # List of CMS screenshots and their captions
                cms_images = [
                    ("media__1784119147468.png", "Antarmuka Halaman Profil Perusahaan pada CMS"),
                    ("media__1784119163487.png", "Dialog Konfirmasi Penyimpanan Perubahan Profil"),
                    ("media__1784119181054.png", "Notifikasi Sukses Penyimpanan Perubahan Profil"),
                    ("media__1784119198342.png", "Antarmuka Halaman Manajemen Struktur Organisasi"),
                    ("media__1784119240845.png", "Antarmuka Halaman Manajemen Bidang Usaha"),
                    ("media__1784119255529.png", "Dialog Tambah Data Bidang Usaha"),
                    ("media__1784119270652.png", "Dialog Edit Data Bidang Usaha"),
                    ("media__1784119305102.png", "Antarmuka Halaman Manajemen Portofolio Proyek"),
                    ("media__1784119317109.png", "Dialog Tambah Data Portofolio Proyek"),
                    ("media__1784119343849.png", "Antarmuka Halaman Galeri Proyek pada CMS"),
                    ("media__1784119385211.png", "Dialog Tambah Foto Baru pada Galeri"),
                    ("media__1784119399617.png", "Antarmuka Halaman Manajemen Divisi & Anggota"),
                    ("media__1784119410622.png", "Dialog Buat Divisi Baru")
                ]
                
                for filename, caption_title in cms_images:
                    img_counters[ch_num] += 1
                    path = f"{base_dir}\\{filename}"
                    caption = f"Gambar {ch_num}.{img_counters[ch_num]} {caption_title}"
                    add_image_from_file(doc, path, width_cm=14.0, caption_text=caption)
            except Exception as e:
                print(f"  Dashboard/CMS screenshots err: {e}")
            i += 1; continue

        if '*(Masukkan Screenshot Halaman Landing Page)*' in stripped:
            try:
                ch_num = 4
                base_dir = r"C:\Users\Haikal Al-Ghifari\.gemini\antigravity-ide\brain\4987b58e-f65c-4ba7-9cfe-908757d8bafe"
                
                img_counters[ch_num] += 1
                caption_id = f"Gambar {ch_num}.{img_counters[ch_num]} Antarmuka Halaman Landing Page (Bahasa Indonesia)"
                new_landing_id_path = f"{base_dir}\\media__1784119075781.png"
                add_image_from_file(doc, new_landing_id_path, width_cm=14.0, caption_text=caption_id)
                
                img_counters[ch_num] += 1
                caption_en = f"Gambar {ch_num}.{img_counters[ch_num]} Antarmuka Halaman Landing Page (Bahasa Inggris)"
                add_image_from_file(doc, SCREENSHOT_LANDING_EN_FILE, width_cm=14.0, caption_text=caption_en)
            except Exception as e:
                print(f"  Landing screenshot err: {e}")
            i += 1; continue

        if 'ss_usecase_diagram.png' in stripped:
            try:
                base_dir = r"C:\Users\Haikal Al-Ghifari\.gemini\antigravity-ide\brain\4987b58e-f65c-4ba7-9cfe-908757d8bafe"
                path = f"{base_dir}\\ss_usecase_diagram.png"
                add_image_from_file(doc, path, width_cm=14.0)
            except Exception as e:
                print(f"  Use Case diagram err: {e}")
            i += 1; continue
        
        # Mermaid
        mm = re.match(r'\[\[MERMAID:(\d+)\]\]', stripped)
        if mm:
            idx = int(mm.group(1))
            img = mermaid_images[idx] if idx < len(mermaid_images) else None
            if img:
                is_activity_diag = (last_level == 3) and last_text.strip().startswith('3.4.')
                is_sequence_diag = (last_level == 3) and last_text.strip().startswith('3.5.')
                is_erd_diag      = (last_level == 2) and last_text.strip().startswith('3.6')
                
                if is_activity_diag:
                    m_h = 11.5   # Activity diagram fits comfortably with its caption on a single page
                    m_w = MAX_IMG_W_CM
                elif is_sequence_diag:
                    m_h = 7.5    # Sequence diagram fits 2 per page comfortably without overflow
                    m_w = MAX_IMG_W_CM
                elif is_erd_diag:
                    m_h = 19.5   # ERD diagram occupies the full portrait page (very tall)
                    m_w = 16.0   # 16.0 cm width in portrait (narrow margins)
                else:
                    m_h = 7.5
                    m_w = MAX_IMG_W_CM
                    
                # Determine figure caption text dynamically using last_text
                m_cap = re.match(r'^([\d]+)\.[\d\.]*\s+(.+)$', last_text.strip())
                if m_cap:
                    ch_num = int(m_cap.group(1))
                    diagram_name = m_cap.group(2)
                else:
                    ch_num = current_chapter
                    diagram_name = last_text
                
                img_counters[ch_num] += 1
                caption = f"Gambar {ch_num}.{img_counters[ch_num]} {diagram_name}"
                
                add_image_from_pil(doc, img, max_width_cm=m_w, max_height_cm=m_h, caption_text=caption)
            i += 1; continue
        
        # Heading (Rule 10)
        # h1 → 14pt bold centered (BAB title or section title)
        # h2+ → 12pt bold LEFT-justified (Subbab)
        hm = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if hm:
            level = len(hm.group(1))
            text  = hm.group(2)
            
            # Store heading context
            last_level = level
            last_text  = text
            
            # Force page break / orientation change
            # Exempt 3.4.1 (Login Admin) so it stays on the intro page under 3.4
            is_activity_subheading = (level == 3) and text.strip().startswith('3.4.') and not text.strip().startswith('3.4.1')
            is_sequence_subheading = False
            if level == 3 and text.strip().startswith('3.5.'):
                m_seq = re.match(r'^3\.5\.(\d+)', text.strip())
                if m_seq:
                    seq_num = int(m_seq.group(1))
                    if seq_num >= 3 and seq_num % 2 == 1:
                        is_sequence_subheading = True
            is_erd_heading = (level == 2) and text.strip().startswith('3.6')
            
            if is_erd_heading:
                from docx.enum.section import WD_ORIENT, WD_SECTION
                erd_sec = doc.add_section(WD_SECTION.NEW_PAGE)
                erd_sec.orientation = WD_ORIENT.PORTRAIT
                w, h = erd_sec.page_width, erd_sec.page_height
                erd_sec.page_width = min(w, h)
                erd_sec.page_height = max(w, h)
                # Narrow margins for portrait ERD page
                erd_sec.top_margin = Cm(2.5)
                erd_sec.bottom_margin = Cm(2.5)
                erd_sec.left_margin = Cm(2.5)
                erd_sec.right_margin = Cm(2.5)
            elif is_activity_subheading or is_sequence_subheading:
                page_break(doc)
                
            if level == 1:
                # BAB title: 14pt, centered, bold
                add_heading(doc, text, level=level, center=True, size=14)
            else:
                # Subbab: 12pt, left-aligned, bold
                add_heading(doc, text, level=level, center=False, size=12)
            ul_letter_idx = 0   # reset ul counter on heading
            i += 1; continue
        
        # HTML table
        if re.match(r'<table', stripped, re.IGNORECASE):
            tbl_lines = []
            while i < len(lines) and not re.search(r'</table>', lines[i], re.IGNORECASE):
                tbl_lines.append(lines[i])
                i += 1
            if i < len(lines):
                tbl_lines.append(lines[i]); i += 1
            raw = '\n'.join(tbl_lines)
            rows_html = re.findall(r'<tr[^>]*>(.*?)</tr>', raw, re.DOTALL|re.IGNORECASE)
            if rows_html:
                header, data_rows = [], []
                for ri, rh in enumerate(rows_html):
                    th_cells = re.findall(r'<th[^>]*>(.*?)</th>', rh, re.DOTALL|re.IGNORECASE)
                    td_cells = re.findall(r'<td[^>]*>(.*?)</td>', rh, re.DOTALL|re.IGNORECASE)
                    if th_cells:
                        if not header: header = [clean_text(c) for c in th_cells]
                        else: data_rows.append([clean_text(c) for c in th_cells])
                    elif td_cells:
                        data_rows.append([clean_text(c) for c in td_cells])
                if not header and data_rows:
                    header = data_rows.pop(0)
                if header:
                    # Check if this is a layout/signature table (like in Kata Pengantar)
                    is_layout = ('Mochammad Alif' in raw) or ('border="0"' in raw)
                    if is_layout:
                        # Render borderless table, no caption, no counter increment
                        add_table_from_md(doc, header, data_rows, borderless=True)
                        doc.add_paragraph()
                    else:
                        tbl_counters[current_chapter] += 1
                        tbl_name = re.sub(r'^[\d\wA-Z\.]+\s*', '', last_text)
                        caption = f"Tabel {current_chapter}.{tbl_counters[current_chapter]} {tbl_name}"
                        
                        add_caption(doc, caption, is_table=True)
                        
                        add_table_from_md(doc, header, data_rows)
                        doc.add_paragraph()
            continue
        
        # Caption detection: "Tabel X.Y" or "Gambar X.Y" style bold lines
        CAP_RE = re.compile(r'^\*{0,2}(Tabel|Gambar|Diagram|Grafik|Bagan)\s+[\d.]+', re.IGNORECASE)
        if CAP_RE.match(stripped):
            # Look ahead: if NEXT non-empty line is a table → this is a TABLE caption (above)
            # Otherwise it's an IMAGE caption (below image already rendered)
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            next_line = lines[j].strip() if j < len(lines) else ''
            is_tbl_cap = next_line.startswith('|')
            add_caption(doc, stripped, is_table=is_tbl_cap)
            i += 1; continue

        # Markdown table – check for pending caption (already added before reaching here)
        if stripped.startswith('|'):
            tbl_buf = [stripped]; i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_buf.append(lines[i].strip()); i += 1
            hdr, drows = None, []
            for tl in tbl_buf:
                if re.match(r'^\|?[\-: |]+\|?$', tl): continue
                cells = [c.strip() for c in re.split(r'(?<!\\)\|', tl.strip('|'))]
                if hdr is None: hdr = cells
                else: drows.append(cells)
            if hdr:
                tbl_counters[current_chapter] += 1
                tbl_name = re.sub(r'^[\d\wA-Z\.]+\s*', '', last_text)
                caption = f"Tabel {current_chapter}.{tbl_counters[current_chapter]} {tbl_name}"
                
                add_caption(doc, caption, is_table=True)
                
                add_table_from_md(doc, hdr, drows)
                doc.add_paragraph()
            continue
        
        # Ordered list – resets ul letter counter
        ol_m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if ol_m:
            ul_letter_idx = 0
            add_list_item(doc, ol_m.group(2), ordered=True, num=int(ol_m.group(1)))
            i += 1; continue

        # Unordered list – use letters a. b. c. ... (Rule 4)
        ul_m = re.match(r'^[-*+]\s+(.+)$', stripped)
        if ul_m:
            ul_letter_idx += 1
            letter = chr(ord('a') + (ul_letter_idx - 1) % 26)
            add_list_item(doc, ul_m.group(1), ordered=False, letter=letter)
            i += 1; continue

        # Non-list line resets ul letter counter
        ul_letter_idx = 0
        
        # Raw HTML tags
        if re.match(r'^<[a-zA-Z/!]', stripped):
            i += 1; continue
        
        # Check if this is a "Sistematika BAB" title, e.g. "**BAB 1: PENDAHULUAN**"
        bab_sistematika_m = re.match(r'^\*\*BAB\s+(\d+)(?:[:\s]+)(.+?)\*\*$', stripped, re.IGNORECASE)
        if bab_sistematika_m:
            num = bab_sistematika_m.group(1)
            title = bab_sistematika_m.group(2).strip()
            # Left aligned, bold, size 12, no indent
            add_para(doc, f"BAB {num} {title.upper()}", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=12, space_before=12, space_after=3, first_indent=False)
            # The next paragraph description SHOULD have the regular first-line indent
            next_para_no_indent = False
            i += 1; continue

        # English Keywords: bold+italic label, italic keywords, no indent
        if stripped.lower().startswith('**keywords:') or stripped.lower().startswith('keywords:'):
            plain = clean_text(stripped)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.first_line_indent = Cm(0)
            
            if ':' in plain:
                lbl, val = plain.split(':', 1)
                r1 = p.add_run(lbl + ':')
                set_font(r1, size=12, bold=True, italic=True)
                r2 = p.add_run(val)
                set_font(r2, size=12, bold=False, italic=True)
            else:
                r = p.add_run(plain)
                set_font(r, size=12, bold=False, italic=True)
            i += 1; continue

        # Indonesian Kata Kunci: bold label, normal keywords, no indent
        if stripped.lower().startswith('**kata kunci:') or stripped.lower().startswith('kata kunci:'):
            plain = clean_text(stripped)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.first_line_indent = Cm(0)
            
            if ':' in plain:
                lbl, val = plain.split(':', 1)
                r1 = p.add_run(lbl + ':')
                set_font(r1, size=12, bold=True, italic=False)
                r2 = p.add_run(val)
                set_font(r2, size=12, bold=False, italic=False)
            else:
                r = p.add_run(plain)
                set_font(r, size=12, bold=False, italic=False)
            i += 1; continue

        # Bibliography entry, e.g. "[1] D. Eka..."
        bib_m = re.match(r'^\[(\d+)\]\s+(.+)$', stripped)
        if bib_m:
            num = bib_m.group(1)
            content = bib_m.group(2)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.first_line_indent = Cm(-0.75)  # Hanging indent
            p.paragraph_format.left_indent = Cm(0.75)
            
            # Add prefix [1]
            r_num = p.add_run(f"[{num}]\t")
            set_font(r_num, size=12)
            
            # Parse inline markdown formatting (bold, italic)
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    set_font(r, size=12, bold=True)
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    set_font(r, size=12, italic=True)
                else:
                    clean = clean_text(part)
                    if clean:
                        r = p.add_run(clean)
                        set_font(r, size=12)
            i += 1; continue

        # Regular paragraph — always JUSTIFY except on cover page
        plain = clean_text(stripped)
        if plain:
            al = WD_ALIGN_PARAGRAPH.CENTER if center_body else WD_ALIGN_PARAGRAPH.JUSTIFY
            indent = (not center_body) and (not next_para_no_indent)
            is_abstract_sec = (sec_name == 'abstract')
            add_para(doc, plain, align=al, italic=is_abstract_sec, first_indent=indent)
            next_para_no_indent = False
        i += 1

# ─────────────────────────────────────────────────────────────
# Page numbers (Rules 8 & 9) – arabic numerals in footer
# ─────────────────────────────────────────────────────────────
from docx.oxml.ns import qn as _qn
from docx.oxml    import OxmlElement as _OE

def add_page_number_footer(section, fmt_switch=None):
    """Add centered page number in footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear existing footer
    for p in footer.paragraphs:
        p.clear()
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # Insert field for page number
    fldChar1 = _OE('w:fldChar'); fldChar1.set(_qn('w:fldCharType'), 'begin')
    instrText = _OE('w:instrText')
    if fmt_switch:
        instrText.text = f'PAGE \\* {fmt_switch}'
    else:
        instrText.text = 'PAGE'
    fldChar2 = _OE('w:fldChar'); fldChar2.set(_qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# Force Word to update fields (including Table of Contents and page numbers) on open
doc.settings.update_fields = False

# Apply page numbers to all sections EXCEPT Cover (Section 0)
for sec_idx, sec in enumerate(doc.sections):
    if sec_idx == 0:
        sec.footer.is_linked_to_previous = False
        for p in sec.footer.paragraphs:
            p.clear()
    elif sec_idx == 1:
        # Preface section (Lembar Pengesahan to Daftar Tabel) uses Roman numerals
        add_page_number_footer(sec, fmt_switch='roman')
    else:
        # Chapters sections use Arabic numerals
        add_page_number_footer(sec, fmt_switch='arabic')

# ─────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"\nDone! Saved to {OUT_FILE}")

# ─────────────────────────────────────────────────────────────
# Convert to PDF using Word COM automation
# ─────────────────────────────────────────────────────────────
try:
    import win32com.client
    import os
    docx_abs = os.path.abspath(OUT_FILE)
    pdf_abs = docx_abs.replace('.docx', '.pdf')
    print(f"Converting DOCX to PDF using Word COM...")
    
    word = None
    doc_word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone - disable all dialog boxes!
        doc_word = word.Documents.Open(docx_abs, ConfirmConversions=False, ReadOnly=False)
        
        # Force update of all fields including TOC, TOF, TOT
        try:
            for toc in doc_word.TablesOfContents:
                toc.Update()
            for tof in doc_word.TablesOfFigures:
                tof.Update()
            doc_word.Save()
            print("Fields updated and DOCX saved.")
        except Exception as update_err:
            print(f"Warning: Field update issue: {update_err}")

        doc_word.SaveAs(pdf_abs, FileFormat=17) # 17 is wdFormatPDF
        print(f"PDF generated successfully at {pdf_abs}")
    finally:
        if doc_word:
            try:
                doc_word.Close(False) # wdDoNotSaveChanges
            except:
                pass
        if word:
            try:
                word.Quit()
            except:
                pass
except Exception as e:
    print(f"Error converting to PDF: {e}")
