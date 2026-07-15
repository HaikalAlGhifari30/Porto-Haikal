import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def set_font(run, size=12, bold=False, italic=False, underline=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def clean_text(text):
    return text.strip()

def add_caption(doc, text, is_table=False):
    p = doc.add_paragraph(style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_table:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
    else:
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.keep_with_next = True
    
    clean = clean_text(text)
    r = p.add_run(clean)
    set_font(r, size=12, bold=True)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p

doc = docx.Document()
p = add_caption(doc, "Gambar 1.1 Test Caption", is_table=False)
print("Paragraph Style Name:", p.style.name)
print("Paragraph XML:", p._p.xml)
